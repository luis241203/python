import tkinter, tkinter.filedialog, re, tkinter.messagebox 
from PIL import Image
root = tkinter.Tk() 
root.withdraw() 
tkinter.messagebox.showinfo(message='Selecciona tu imagen', title='AVISO')
file_path = tkinter.filedialog.askopenfilename() 
match = re.search(r'/.*\..+', file_path)
file_position = file_path.find(match.group())
Imagen=Image.open(file_path)
new_image=Imagen.resize((132,188))
new_image.save('/home/luis/Escritorio/fotos/photorecortada.png')

from docx import Document
from docx.shared import Inches
documento=Document()
documento.add_heading('FOTOS',0)
documento.add_picture('photorecortada.png', width=Inches(1.37795276), height=Inches(1.96850394))
documento.add_picture('photorecortada.png', width=Inches(1.37795276), height=Inches(1.96850394))
documento.add_picture('photorecortada.png', width=Inches(1.37795276), height=Inches(1.96850394))
documento.add_picture('photorecortada.png', width=Inches(1.37795276), height=Inches(1.96850394))
documento.save('/home/luis/Escritorio/fotos/photo.docx')

tkinter.messagebox.showinfo(message='SU DOCUMENTO SE ENCUENTRA EN LA SIGUIENTE RUTA: /home/luis/Escritorio/fotos', title='¡PR0CESO EXITOSO!')